#!/usr/bin/env python3
"""
Essay Grader
============
Reads .docx essays from a folder, scores them against a rubric using Claude AI,
and outputs a formatted Excel gradebook.

Usage:
    python grade_essays.py --essays ./essays --rubric rubric.txt --output grades.xlsx
    python grade_essays.py --essays ./essays --rubric rubric.txt  # output defaults to grades.xlsx
"""

import json
import os
import sys
import time
from pathlib import Path

import anthropic
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

# ── Colors ────────────────────────────────────────────────────────────────────
HEADER_BG   = "1F4E79"
HEADER_FG   = "FFFFFF"
ALT_ROW     = "D6E4F0"
WHITE       = "FFFFFF"
C_GREEN     = "C6EFCE"   # score 4
C_YELLOW    = "FFEB9C"   # score 3
C_ORANGE    = "FFCC99"   # score 2
C_RED       = "FFC7CE"   # score 1
C_GRAY      = "F2F2F2"   # missing / n/a
STATUS_COLORS = {
    "Complete":   "C6EFCE",
    "Incomplete": "FFEB9C",
    "Minimal":    "FFC7CE",
    "Blank":      "FFC7CE",
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def read_docx(path: Path) -> str:
    """Extract plain text from a .docx file."""
    try:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception as e:
        return f"[ERROR reading file: {e}]"


def extract_student_name(filepath: Path) -> str:
    """
    Extract the student's name from the MLA header of the essay (first non-empty line).
    Handles variants like 'Name: Sara Stone' or '89Rikki shaul' (leading junk).
    Falls back to the filename if the docx can't be read or the name looks wrong.
    """
    import re

    try:
        doc = Document(filepath)
        first_line = next(
            (p.text.strip() for p in doc.paragraphs if p.text.strip()), ""
        )

        if first_line:
            # Strip common label prefixes: "Name:", "Name ", etc.
            cleaned = re.sub(r'^name\s*[:\-]?\s*', '', first_line, flags=re.IGNORECASE).strip()

            # Strip leading non-letter junk (e.g. "89Rikki" → "Rikki")
            cleaned = re.sub(r'^[^a-zA-Z\'\"]+', '', cleaned).strip()

            # Accept if it looks like a real name: 2–5 words, no more than 40 chars,
            # doesn't look like a class/teacher line
            words = cleaned.split()
            is_plausible = (
                2 <= len(words) <= 5
                and len(cleaned) <= 40
                and not re.search(r'(period|class|teacher|burkinsky|ela\b)', cleaned, re.IGNORECASE)
            )
            if is_plausible:
                return cleaned
    except Exception:
        pass

    # Fallback: derive from filename
    stem = filepath.stem.lstrip("_").strip()
    for sep in [" - ", " – ", "-Essay", "_-_"]:
        if sep in stem:
            stem = stem.split(sep)[0]
            break
    return stem.replace("_", " ").strip()


def score_color(score) -> str:
    if score is None: return C_GRAY
    return {4: C_GREEN, 3: C_YELLOW, 2: C_ORANGE, 1: C_RED}.get(score, WHITE)


def letter_grade(total, max_pts) -> str:
    if total is None or max_pts == 0: return "N/A"
    pct = total / max_pts * 100
    if pct >= 93: return "A"
    if pct >= 90: return "A-"
    if pct >= 87: return "B+"
    if pct >= 83: return "B"
    if pct >= 80: return "B-"
    if pct >= 77: return "C+"
    if pct >= 73: return "C"
    if pct >= 70: return "C-"
    if pct >= 65: return "D"
    return "F"


# ── AI Grading ────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an experienced English teacher grading student essays.
You will receive a rubric and an essay. Score the essay according to the rubric.

You MUST respond with valid JSON only — no markdown fences, no extra text.

Return exactly this structure:
{
  "status": "Complete" | "Incomplete" | "Minimal" | "Blank",
  "categories": {
    "<category_name>": <score_1_to_4_or_null>
  },
  "notes": "<concise teacher feedback, 1-3 sentences>"
}

Rules:
- Use null for a category if that section is entirely absent from the essay.
- "Blank" = file is empty or only contains a template with no student writing.
- "Minimal" = less than one full paragraph of real content.
- "Incomplete" = has real content but is missing one or more required sections.
- "Complete" = all required sections are present.
- Keep notes brief and actionable (what's strong, what's missing, key errors).
- Score integers only: 1, 2, 3, or 4.
"""


def load_examples(examples_dir: Path) -> list[dict]:
    """
    Load paired example essays + teacher-filled rubrics from examples_dir.

    Naming convention:
      sara.docx          <- sample essay
      sara_rubric.docx   <- your filled-in scores/notes for that essay

    Returns a list of {"essay": str, "rubric": str} dicts.
    """
    examples = []
    essay_files = [
        f for f in sorted(examples_dir.glob("*.docx"))
        if not f.stem.endswith("_rubric")
    ]
    for essay_path in essay_files:
        rubric_path = examples_dir / f"{essay_path.stem}_rubric.docx"
        if not rubric_path.exists():
            print(f"  ⚠ No matching rubric found for '{essay_path.name}' — skipping")
            continue
        essay_text  = read_docx(essay_path)
        rubric_text = read_docx(rubric_path)
        if essay_text and rubric_text:
            examples.append({"essay": essay_text, "rubric": rubric_text})
            print(f"  ✓ Example loaded: {essay_path.stem}")
    return examples


def grade_with_ai(client: anthropic.Anthropic, rubric_text: str, essay_text: str,
                  student_name: str, category_names: list[str],
                  examples: list[dict] = None) -> dict:
    """Call Claude to score one essay. Returns parsed JSON dict."""

    if not essay_text or len(essay_text) < 30:
        return {
            "status": "Blank",
            "categories": {c: None for c in category_names},
            "notes": "File was empty or contained only a template."
        }

    # Build an explicit JSON template showing exactly what keys to fill in
    categories_template = ",\n    ".join(
        f'"{c}": <1|2|3|4|null>' for c in category_names
    )

    def build_user_msg(essay, name):
        return f"""RUBRIC:
{rubric_text}

---

STUDENT: {name}
ESSAY:
{essay}

---

Grade this essay against every category in the rubric.
You MUST score ALL of the following categories — do not skip any:
{chr(10).join(f'  - {c}' for c in category_names)}

Return this exact JSON (fill in every value, use null only if the section is completely absent):
{{{{
  "status": "Complete" | "Incomplete" | "Minimal" | "Blank",
  "categories": {{{{
    {categories_template}
  }}}},
  "notes": "<concise feedback, 1-3 sentences>"
}}}}"""

    # Build conversation — prepend few-shot examples as prior turns if provided
    messages = []
    for ex in (examples or []):
        messages.append({"role": "user",      "content": build_user_msg(ex["essay"], "Example Student")})
        messages.append({"role": "assistant", "content": ex["rubric"]})
    messages.append({"role": "user", "content": build_user_msg(essay_text, student_name)})

    for attempt in range(3):
        try:
            response = client.messages.create(
                model="claude-opus-4-6",
                max_tokens=1000,
                system=SYSTEM_PROMPT,
                messages=messages
            )
            raw = response.content[0].text.strip()
            # Strip accidental markdown fences
            if raw.startswith("```"):
                raw = raw.split("```")[1]
                if raw.startswith("json"):
                    raw = raw[4:]
            return json.loads(raw.strip())
        except json.JSONDecodeError as e:
            print(f"  ⚠ JSON parse error (attempt {attempt+1}): {e}")
            time.sleep(1)
        except anthropic.RateLimitError:
            wait = 20 * (attempt + 1)
            print(f"  ⚠ Rate limit hit — waiting {wait}s...")
            time.sleep(wait)
        except Exception as e:
            print(f"  ⚠ API error (attempt {attempt+1}): {e}")
            time.sleep(2)

    return {
        "status": "Error",
        "categories": {},
        "notes": "Grading failed after 3 attempts — please grade manually."
    }


# ── Excel Output ──────────────────────────────────────────────────────────────

def build_excel(results: list[dict], category_names: list[str], output_path: Path,
                essay_title: str = "", period: str = "", grade_level: str = "",
                weights: dict = None, rubric_total: int = None):
    wb = Workbook()
    ws = wb.active
    ws.title = "Essay Grades"

    if weights is None:
        weights = {c: 4 for c in category_names}
    if rubric_total is None:
        rubric_total = sum(weights.get(c, 4) for c in category_names)

    # ── Column layout ──────────────────────────────────────────────────────
    # Col 1: Student name
    # Col 2: Status
    # Col 3..N: one per rubric category
    # Col N+1: Total
    # Col N+2: Grade
    # Col N+3: Notes

    n_cats = len(category_names)
    total_col = 3 + n_cats
    grade_col = total_col + 1
    notes_col = grade_col + 1
    last_col  = notes_col

    # ── Title rows ─────────────────────────────────────────────────────────
    # Row 1: Essay title (merged, if provided)
    # Row 2: Subtitle combining grade level + period (merged, if either provided)
    # Then column headers

    title_rows = 0

    if essay_title:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=last_col)
        t = ws.cell(row=1, column=1, value=essay_title)
        t.font = Font(bold=True, size=16, color=HEADER_FG)
        t.fill = PatternFill("solid", start_color=HEADER_BG)
        t.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 32
        title_rows += 1

    # Combine grade level and period into one subtitle row
    subtitle_parts = []
    if grade_level:
        subtitle_parts.append(grade_level)
    if period:
        subtitle_parts.append(f"Period {period}")
    subtitle = "  |  ".join(subtitle_parts)

    if subtitle:
        subtitle_row = title_rows + 1
        ws.merge_cells(start_row=subtitle_row, start_column=1,
                       end_row=subtitle_row, end_column=last_col)
        p = ws.cell(row=subtitle_row, column=1, value=subtitle)
        p.font = Font(bold=True, size=12, color=HEADER_FG)
        p.fill = PatternFill("solid", start_color="2E75B6")
        p.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[subtitle_row].height = 24
        title_rows += 1

    # Column header row sits just below the title block
    header_row = title_rows + 1

    # ── Header row ─────────────────────────────────────────────────────────
    headers = (
        ["Student", "Status"]
        + [f"{c}\n({weights.get(c, 4)})" for c in category_names]
        + [f"TOTAL\n(/{rubric_total})", "Grade", "Notes"]
    )
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col, value=h)
        cell.font = Font(bold=True, color=HEADER_FG, size=10)
        cell.fill = PatternFill("solid", start_color=HEADER_BG)
        cell.alignment = Alignment(wrap_text=True, horizontal="center", vertical="center")
    ws.row_dimensions[header_row].height = 42

    # Freeze below the title + header block
    ws.freeze_panes = f"A{header_row + 1}"

    # ── Data rows ──────────────────────────────────────────────────────────
    for i, r in enumerate(results):
        row = header_row + 1 + i
        bg = ALT_ROW if i % 2 == 0 else WHITE

        scores = [r["categories"].get(c) for c in category_names]

        # Each category is scored 1–4. Total = (sum of scores / max possible) * rubric_total.
        # Missing sections (None) count as 0. Denominator is always rubric_total.
        n_cats      = len(category_names)
        earned_raw  = sum(sc for sc in scores if sc is not None)   # sum of 1–4 scores
        max_raw     = n_cats * 4                                    # e.g. 8 × 4 = 32
        total       = round(earned_raw / max_raw * rubric_total, 1) if max_raw else 0
        # Show as integer when whole (e.g. 24.0 → 24)
        total_display = int(total) if total == int(total) else total

        def cell(col, value, bold=False, fg="000000", bg_=None, center=False, wrap=False):
            c = ws.cell(row=row, column=col, value=value)
            c.font = Font(bold=bold, size=10, color=fg)
            c.fill = PatternFill("solid", start_color=bg_ or bg)
            c.alignment = Alignment(horizontal="center" if center else "left",
                                    vertical="center", wrap_text=wrap)
            return c

        # Name
        cell(1, r["name"], bold=True)

        # Status
        status = r.get("status", "Unknown")
        sc = ws.cell(row=row, column=2, value=status)
        sc.font = Font(bold=True, size=9)
        sc.fill = PatternFill("solid", start_color=STATUS_COLORS.get(status, C_GRAY))
        sc.alignment = Alignment(horizontal="center", vertical="center")

        # Score columns
        for j, sc_val in enumerate(scores):
            col = 3 + j
            c = ws.cell(row=row, column=col)
            c.value = sc_val if sc_val is not None else "—"
            c.font = Font(size=10, bold=(sc_val is not None))
            c.fill = PatternFill("solid", start_color=score_color(sc_val))
            c.alignment = Alignment(horizontal="center", vertical="center")

        # Total — always out of rubric_total; show as integer if whole number
        c = ws.cell(row=row, column=total_col)
        if any(sc is not None for sc in scores):
            c.value = f"{total_display}/{rubric_total}"
        else:
            c.value = "N/A"
        c.font = Font(bold=True, size=10)
        c.fill = PatternFill("solid", start_color=bg)
        c.alignment = Alignment(horizontal="center", vertical="center")

        # Grade
        c = ws.cell(row=row, column=grade_col)
        c.value = letter_grade(total, rubric_total) if any(sc is not None for sc in scores) else "N/A"
        c.font = Font(bold=True, size=10)
        c.fill = PatternFill("solid", start_color=bg)
        c.alignment = Alignment(horizontal="center", vertical="center")

        # Notes
        c = ws.cell(row=row, column=notes_col)
        c.value = r.get("notes", "")
        c.font = Font(size=9, italic=True)
        c.fill = PatternFill("solid", start_color=bg)
        c.alignment = Alignment(wrap_text=True, vertical="center")

        ws.row_dimensions[row].height = 45

    # ── Column widths ──────────────────────────────────────────────────────
    ws.column_dimensions["A"].width = 22  # name
    ws.column_dimensions["B"].width = 12  # status
    for j in range(n_cats):
        col_letter = get_column_letter(3 + j)
        ws.column_dimensions[col_letter].width = 11
    ws.column_dimensions[get_column_letter(total_col)].width = 10
    ws.column_dimensions[get_column_letter(grade_col)].width = 8
    ws.column_dimensions[get_column_letter(notes_col)].width = 55

    # ── Legend sheet ───────────────────────────────────────────────────────
    ls = wb.create_sheet("Legend")
    legend_items = [
        ("Score 4 — Excellent",  C_GREEN),
        ("Score 3 — Good",       C_YELLOW),
        ("Score 2 — Developing", C_ORANGE),
        ("Score 1 — Beginning",  C_RED),
        ("— Missing / not submitted", C_GRAY),
    ]
    ls.cell(1, 1, "COLOR KEY").font = Font(bold=True, size=12)
    for i, (label, color) in enumerate(legend_items, 2):
        c = ls.cell(i, 1, label)
        c.fill = PatternFill("solid", start_color=color)
        c.font = Font(size=10)
    ls.column_dimensions["A"].width = 30

    wb.save(output_path)
    print(f"\n✅  Saved: {output_path}")


# ── Main ──────────────────────────────────────────────────────────────────────

def parse_rubric_metadata(rubric_text: str, rubric_path: Path = None) -> tuple[str, list[str], dict[str, int], int]:
    """
    Extract title, category names, per-category weights, and the rubric's stated total.

    If rubric_path is a .docx, reads the table directly — reliable and exact.
    Otherwise asks Claude to parse the plain text.

    Returns:
        title          (str)
        category_names (list[str])
        weights        (dict: name → points)
        rubric_total   (int)  ← the number found next to "TOTAL SCORE" in the rubric
    """

    # ── Strategy 1: parse the .docx table directly ────────────────────────
    if rubric_path and rubric_path.suffix.lower() == ".docx":
        try:
            return _parse_rubric_from_docx(rubric_path)
        except Exception as e:
            print(f"  ⚠ Direct table parse failed ({e}), falling back to AI...")

    # ── Strategy 2: ask Claude (for .txt rubrics) ─────────────────────────
    return _parse_rubric_with_ai(rubric_text)


def _parse_rubric_from_docx(rubric_path: Path) -> tuple[str, list[str], dict[str, int], int]:
    """Parse a rubric .docx by reading its table rows directly."""
    import re
    doc = Document(rubric_path)

    # Title = first non-empty paragraph
    title = ""
    for p in doc.paragraphs:
        if p.text.strip():
            title = p.text.strip()
            # Strip common prefixes like "Essay Rubric: "
            title = re.sub(r'^(essay\s+)?rubric\s*:\s*', '', title, flags=re.IGNORECASE).strip()
            break

    categories = []
    weights    = {}
    rubric_total = None

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells or not cells[0]:
                continue

            first = cells[0].replace("\n", " ").strip()

            # Detect the TOTAL SCORE row
            if re.search(r'total\s*score', first, re.IGNORECASE):
                # Look for a number like "/ 32" or "32" anywhere in the row
                for cell_text in cells:
                    m = re.search(r'/?\s*(\d+)', cell_text)
                    if m:
                        rubric_total = int(m.group(1))
                        break
                continue

            # Skip header row
            if first.lower() in ("category", ""):
                continue

            # Each data row is one gradeable category
            cat_name = first.replace("\n", " ")
            categories.append(cat_name)
            weights[cat_name] = 4   # every row is scored /4; total comes from rubric_total

    # If the rubric didn't state a total, calculate it
    if rubric_total is None:
        rubric_total = len(categories) * 4

    return title, categories, weights, rubric_total


def _parse_rubric_with_ai(rubric_text: str) -> tuple[str, list[str], dict[str, int], int]:
    """Fallback: use Claude to parse a plain-text rubric."""
    import re
    client = anthropic.Anthropic()
    try:
        resp = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=600,
            messages=[{
                "role": "user",
                "content": (
                    "From this rubric extract:\n"
                    "1. The assignment title (without any 'Rubric:' prefix).\n"
                    "2. Every graded category name, in order. List each body paragraph separately.\n"
                    "3. The point value for each category (each row is scored /4).\n"
                    "4. The total score stated in the rubric (e.g. '/ 32' → 32).\n\n"
                    "Return JSON only:\n"
                    '{"title": "...", "categories": [...], "weights": {...}, "total": 32}\n\n'
                    f"RUBRIC:\n{rubric_text}"
                )
            }]
        )
        raw = resp.content[0].text.strip().strip("`")
        if raw.startswith("json"): raw = raw[4:]
        data       = json.loads(raw)
        title      = data.get("title", "")
        categories = data.get("categories", [])
        weights    = data.get("weights", {c: 4 for c in categories})
        total      = data.get("total", len(categories) * 4)
        for c in categories:
            weights.setdefault(c, 4)
        return title, categories, weights, total
    except Exception:
        # Hard fallback
        lines = [l.strip() for l in rubric_text.splitlines() if l.strip()]
        title = lines[0] if lines else ""
        import re
        total_match = re.search(r'/\s*(\d+)', rubric_text)
        rubric_total = int(total_match.group(1)) if total_match else 32
        cats = ["Introduction", "Body Paragraph 1", "Body Paragraph 2", "Body Paragraph 3",
                "Conclusion", "Spelling and Grammar", "Formatting", "Transition Words"]
        weights = {c: 4 for c in cats}
        return title, cats, weights, rubric_total


def load_config(config_path: Path) -> dict:
    """
    Parse a simple key = value config file.
    Lines starting with # are comments. Blank lines are ignored.
    Returns a dict of all found keys (lowercased, spaces→underscores).
    """
    config = {}
    for line in config_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if "=" not in line:
            continue
        key, _, value = line.partition("=")
        key = key.strip().lower().replace(" ", "_").replace("-", "_")
        value = value.strip()
        # Strip optional inline comments
        if " #" in value:
            value = value[:value.index(" #")].strip()
        config[key] = value
    return config


def write_score_sheets(results: list[dict], category_names: list[str],
                       rubric_total: int, output_path: Path):
    """
    Write a .txt file with one score block per student, e.g.:

    ─────────────────────────────
    Raizy Rosenberg
    ─────────────────────────────
    Introduction: 3
    Body Paragraph 1: 2
    Body Paragraph 2: 3
    Body Paragraph 3: 3
    Conclusion: —
    Spelling and Grammar: 2
    Formatting: 2
    Transition Words: 3

    Total Score: 18/32
    ─────────────────────────────
    """
    n_cats = len(category_names)
    lines = []

    for r in results:
        name   = r.get("name", "Unknown")
        scores = [r["categories"].get(c) for c in category_names]

        # Calculate total (same formula as build_excel)
        earned_raw = sum(sc for sc in scores if sc is not None)
        max_raw    = n_cats * 4
        total      = round(earned_raw / max_raw * rubric_total, 1) if max_raw else 0
        total_display = int(total) if total == int(total) else total

        divider = "─" * 35
        lines.append(divider)
        lines.append(name)
        lines.append(divider)

        for cat, sc in zip(category_names, scores):
            score_str = str(sc) if sc is not None else "—"
            lines.append(f"{cat}: {score_str}")

        lines.append("")
        lines.append(f"Total Score: {total_display}/{rubric_total}")
        lines.append(divider)
        lines.append("")   # blank line between students

    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"📄  Score sheets saved: {output_path}")


def main():
    # ── Locate config file ────────────────────────────────────────────────
    # Default: config.txt next to this script.
    # Override with: python grade_essays.py my_other_config.txt
    if len(sys.argv) > 1:
        config_path = Path(sys.argv[1])
    else:
        config_path = Path(__file__).parent / "config.txt"

    if not config_path.exists():
        sys.exit(
            f"❌  Config file not found: {config_path}\n"
            f"    Create a config.txt next to grade_essays.py (see README for format)."
        )

    cfg = load_config(config_path)
    print(f"⚙️   Config loaded: {config_path}\n")

    # ── Read settings ─────────────────────────────────────────────────────
    essays_str   = cfg.get("essays_folder", "")
    rubric_str   = cfg.get("rubric_file", "")
    output_str   = cfg.get("output_file", "grades.xlsx")
    examples_str = cfg.get("examples_folder", "")
    period       = cfg.get("period", "")
    grade_level  = cfg.get("grade_level", "")
    api_key      = cfg.get("api_key", "") or os.environ.get("ANTHROPIC_API_KEY", "")

    if not essays_str:
        sys.exit("❌  'essays_folder' is not set in config.txt")
    if not rubric_str:
        sys.exit("❌  'rubric_file' is not set in config.txt")

    # ── Validate paths ────────────────────────────────────────────────────
    essays_dir  = Path(essays_str).expanduser().resolve()
    rubric_path = Path(rubric_str).expanduser().resolve()
    output_path = Path(output_str).expanduser().resolve()

    if not essays_dir.is_dir():
        sys.exit(f"❌  Essays folder not found: {essays_dir}")

    if not rubric_path.exists():
        sys.exit(f"❌  Rubric file not found: {rubric_path}")

    # Create output directory if it doesn't exist
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # ── Read rubric ───────────────────────────────────────────────────────
    if rubric_path.suffix.lower() == ".docx":
        rubric_text = read_docx(rubric_path)
    else:
        rubric_text = rubric_path.read_text(encoding="utf-8")
    print(f"📋  Rubric loaded: {rubric_path.name}")

    # ── Set up API client ─────────────────────────────────────────────────
    if not api_key:
        sys.exit(
            "❌  No API key found.\n"
            "    Add 'api_key = your-key-here' to config.txt,\n"
            "    or set the ANTHROPIC_API_KEY environment variable."
        )
    client = anthropic.Anthropic(api_key=api_key)

    # ── Detect rubric title + categories ──────────────────────────────────
    print("🔍  Detecting rubric title and categories...")
    essay_title, category_names, weights, rubric_total = parse_rubric_metadata(rubric_text, rubric_path)
    if essay_title:
        print(f"    Title:      {essay_title}")
    print(f"    Categories: {', '.join(category_names)}")
    print(f"    Total pts:  {rubric_total}")
    if grade_level:
        print(f"    Grade:      {grade_level}")
    if period:
        print(f"    Period:     {period}")

    # ── Load grading examples (optional) ─────────────────────────────────
    examples = []
    if examples_str:
        examples_dir = Path(examples_str).expanduser().resolve()
        if not examples_dir.is_dir():
            print(f"  ⚠ Examples folder not found: {examples_dir} — grading without examples")
        else:
            print(f"\n📚  Loading grading examples from {examples_dir}...")
            examples = load_examples(examples_dir)
            if examples:
                print(f"    {len(examples)} example(s) will guide the AI's grading style\n")
            else:
                print(f"    No valid examples found — grading without examples\n")

    # ── Find essay files ──────────────────────────────────────────────────
    essay_files = sorted(essays_dir.glob("*.docx"))
    if not essay_files:
        sys.exit(f"❌  No .docx files found in {essays_dir}")
    print(f"\n📂  Found {len(essay_files)} essays in {essays_dir}\n")

    # ── Grade each essay ──────────────────────────────────────────────────
    results = []
    for idx, filepath in enumerate(essay_files, 1):
        name = extract_student_name(filepath)
        print(f"  [{idx:>2}/{len(essay_files)}] Grading {name}...", end=" ", flush=True)

        essay_text = read_docx(filepath)
        result = grade_with_ai(client, rubric_text, essay_text, name, category_names, examples)
        result["name"] = name
        result["file"] = filepath.name

        # Ensure all categories are present in the result
        for cat in category_names:
            if cat not in result.get("categories", {}):
                result.setdefault("categories", {})[cat] = None

        valid = [s for s in result["categories"].values() if s is not None]
        earned_raw = sum(valid)
        total_preview = round(earned_raw / (len(category_names) * 4) * rubric_total, 1) if category_names else 0
        print(f"{result['status']}  ({total_preview}/{rubric_total})")

        results.append(result)
        time.sleep(0.3)  # gentle rate limiting

    # ── Write Excel ───────────────────────────────────────────────────────
    print(f"\n📊  Building Excel gradebook...")
    build_excel(results, category_names, output_path,
                essay_title=essay_title, period=period, grade_level=grade_level,
                weights=weights, rubric_total=rubric_total)

    # ── Write score sheets .txt ───────────────────────────────────────────
    txt_path = output_path.with_suffix(".txt")
    write_score_sheets(results, category_names, rubric_total, txt_path)

    # ── Summary ───────────────────────────────────────────────────────────
    from collections import Counter
    statuses = Counter(r["status"] for r in results)
    print("\n── Summary ──────────────────────────────")
    for status, count in sorted(statuses.items()):
        print(f"   {status:<12} {count} student(s)")
    print(f"─────────────────────────────────────────")
    print(f"   Total: {len(results)} essays graded")


if __name__ == "__main__":
    main()
